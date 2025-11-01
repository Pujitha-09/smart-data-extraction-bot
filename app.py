import streamlit as st
import pandas as pd
from io import BytesIO
from scraper import scrape_data, generate_summary
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from docx import Document

# =========================================================
# PAGE CONFIGURATION
# =========================================================
st.set_page_config(
    page_title="Smart Data Extraction Bot 🧠",
    page_icon="🧠",
    layout="centered",
)

st.title("🧠 Smart Data Extraction Bot")
st.caption("Extract and summarize website content using AI | Built with Python, Streamlit & Hugging Face")
st.divider()


# =========================================================
# UTILITIES: PDF & DOCX EXPORTERS
# =========================================================
def generate_pdf(summary_text):
    """Generate a downloadable PDF from summary text."""
    buffer = BytesIO()
    c = canvas.Canvas(buffer, pagesize=letter)
    width, height = letter

    c.setFont("Helvetica-Bold", 14)
    c.drawString(50, height - 80, "Smart Data Extraction Bot - Summary Report")

    c.setFont("Helvetica", 11)
    text_object = c.beginText(50, height - 120)
    wrap_limit = 95

    for line in summary_text.split("\n"):
        while len(line) > wrap_limit:
            text_object.textLine(line[:wrap_limit])
            line = line[wrap_limit:]
        text_object.textLine(line)

    c.drawText(text_object)
    c.showPage()
    c.save()

    buffer.seek(0)
    return buffer


def generate_docx(summary_text):
    """Generate a downloadable Word DOCX file from summary text."""
    doc = Document()
    doc.add_heading("Smart Data Extraction Bot - Summary Report", level=1)
    doc.add_paragraph(summary_text)
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


# =========================================================
# MAIN INTERFACE
# =========================================================
url = st.text_input("🔗 Enter a Website URL:")

style = st.selectbox(
    "🎨 Choose Summary Style:",
    ["Professional", "Concise"],
    help="Select how you want the summary to sound.",
)

if st.button("🚀 Extract & Summarize"):
    if not url.strip():
        st.warning("⚠️ Please enter a valid URL before proceeding.")
    else:
        with st.spinner("🔍 Extracting content from website..."):
            df = scrape_data(url)

        st.success(f"✅ Extracted {len(df)} item(s) successfully!")
        st.dataframe(df, use_container_width=True)

        with st.spinner("🧠 Generating AI-powered summary..."):
            summary = generate_summary(df, style=style)

        st.divider()
        st.subheader("📋 AI Summary")
        st.write(summary)

        # =========================================================
        # EXPORT OPTIONS
        # =========================================================
        pdf_buffer = generate_pdf(summary)
        docx_buffer = generate_docx(summary)

        st.markdown("### 💾 Download Summary")
        col1, col2 = st.columns(2)

        with col1:
            st.download_button(
                label="📄 Download as PDF",
                data=pdf_buffer,
                file_name="summary_report.pdf",
                mime="application/pdf",
            )

        with col2:
            st.download_button(
                label="📝 Download as DOCX",
                data=docx_buffer,
                file_name="summary_report.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )

        st.divider()
        st.info("✅ Done! You can now download or share your AI-generated summary report.")


# =========================================================
# FOOTER
# =========================================================
st.markdown(
    """
    ---
    👩‍💻 **Developed by [Vari Naga Pujitha](https://www.linkedin.com/in/varinagapujitha)**  
    Powered by [Streamlit](https://streamlit.io) • [Hugging Face Transformers](https://huggingface.co/models)
    """
)
